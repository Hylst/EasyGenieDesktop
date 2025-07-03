#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Easy Genie Desktop - Export Service Manager

Manages data export to various formats (PDF, DOCX, TXT, JSON, CSV).
"""

import logging
import json
import csv
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from datetime import datetime
from enum import Enum
import io

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
except ImportError:
    reportlab = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = None


class ExportFormat(Enum):
    """Export format types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    JSON = "json"
    CSV = "csv"
    HTML = "html"
    MD = "md"  # Markdown


class ExportTemplate(Enum):
    """Export template types."""
    SIMPLE = "simple"
    DETAILED = "detailed"
    PROFESSIONAL = "professional"
    MINIMAL = "minimal"


class ExportServiceManager:
    """Manages data export to various formats."""
    
    def __init__(self, settings_manager=None):
        """Initialize export service manager."""
        self.logger = logging.getLogger(__name__)
        self.settings_manager = settings_manager
        
        # Export settings
        self.default_format = ExportFormat.PDF
        self.default_template = ExportTemplate.SIMPLE
        self.output_directory = Path.home() / "Documents" / "Easy Genie Exports"
        
        # Document settings
        self.page_size = A4
        self.margins = {
            'top': 1 * inch,
            'bottom': 1 * inch,
            'left': 1 * inch,
            'right': 1 * inch
        }
        
        # Style settings
        self.font_family = "Helvetica"
        self.font_size = 12
        self.header_color = colors.HexColor('#2E86AB')
        self.accent_color = colors.HexColor('#A23B72')
        
        # Load settings
        self._load_settings()
        
        # Ensure output directory exists
        self.output_directory.mkdir(parents=True, exist_ok=True)
    
    def _load_settings(self):
        """Load export settings from configuration."""
        if not self.settings_manager:
            return
        
        try:
            # Load export preferences
            format_name = self.settings_manager.get('export.default_format', 'PDF')
            try:
                self.default_format = ExportFormat[format_name]
            except KeyError:
                self.default_format = ExportFormat.PDF
            
            template_name = self.settings_manager.get('export.default_template', 'SIMPLE')
            try:
                self.default_template = ExportTemplate[template_name]
            except KeyError:
                self.default_template = ExportTemplate.SIMPLE
            
            # Load output directory
            output_dir = self.settings_manager.get('export.output_directory', '')
            if output_dir:
                self.output_directory = Path(output_dir)
            
            # Load style settings
            self.font_family = self.settings_manager.get('export.font_family', 'Helvetica')
            self.font_size = self.settings_manager.get('export.font_size', 12)
            
            self.logger.info("Export settings loaded")
            
        except Exception as e:
            self.logger.error(f"Failed to load export settings: {e}")
    
    def export_tasks(self, tasks: List[Dict], format_type: ExportFormat = None, 
                    template: ExportTemplate = None, filename: str = None) -> Optional[Path]:
        """Export tasks to specified format."""
        format_type = format_type or self.default_format
        template = template or self.default_template
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"tasks_export_{timestamp}.{format_type.value}"
        
        output_path = self.output_directory / filename
        
        try:
            if format_type == ExportFormat.PDF:
                return self._export_tasks_pdf(tasks, output_path, template)
            elif format_type == ExportFormat.DOCX:
                return self._export_tasks_docx(tasks, output_path, template)
            elif format_type == ExportFormat.TXT:
                return self._export_tasks_txt(tasks, output_path, template)
            elif format_type == ExportFormat.JSON:
                return self._export_tasks_json(tasks, output_path)
            elif format_type == ExportFormat.CSV:
                return self._export_tasks_csv(tasks, output_path)
            elif format_type == ExportFormat.HTML:
                return self._export_tasks_html(tasks, output_path, template)
            elif format_type == ExportFormat.MD:
                return self._export_tasks_markdown(tasks, output_path, template)
            else:
                raise ValueError(f"Unsupported format: {format_type}")
                
        except Exception as e:
            self.logger.error(f"Failed to export tasks: {e}")
            return None
    
    def _export_tasks_pdf(self, tasks: List[Dict], output_path: Path, 
                         template: ExportTemplate) -> Optional[Path]:
        """Export tasks to PDF format."""
        if not reportlab:
            raise ImportError("reportlab is required for PDF export")
        
        try:
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=self.page_size,
                topMargin=self.margins['top'],
                bottomMargin=self.margins['bottom'],
                leftMargin=self.margins['left'],
                rightMargin=self.margins['right']
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=self.header_color,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=self.header_color,
                spaceBefore=20,
                spaceAfter=10
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=self.font_size,
                spaceBefore=6,
                spaceAfter=6
            )
            
            # Build document content
            story = []
            
            # Title
            title = Paragraph("Rapport des Tâches - Easy Genie", title_style)
            story.append(title)
            
            # Generation info
            gen_info = Paragraph(
                f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}",
                body_style
            )
            story.append(gen_info)
            story.append(Spacer(1, 20))
            
            # Summary
            total_tasks = len(tasks)
            completed_tasks = len([t for t in tasks if t.get('status') == 'completed'])
            pending_tasks = total_tasks - completed_tasks
            
            summary_data = [
                ['Statistiques', ''],
                ['Total des tâches', str(total_tasks)],
                ['Tâches terminées', str(completed_tasks)],
                ['Tâches en cours', str(pending_tasks)]
            ]
            
            summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), self.header_color),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(summary_table)
            story.append(Spacer(1, 30))
            
            # Tasks section
            if template == ExportTemplate.DETAILED:
                story.append(Paragraph("Détail des Tâches", heading_style))
                
                for i, task in enumerate(tasks, 1):
                    # Task header
                    task_title = f"{i}. {task.get('title', 'Sans titre')}"
                    story.append(Paragraph(task_title, heading_style))
                    
                    # Task details
                    details = []
                    if task.get('description'):
                        details.append(f"<b>Description:</b> {task['description']}")
                    
                    details.append(f"<b>Statut:</b> {task.get('status', 'pending')}")
                    details.append(f"<b>Priorité:</b> {task.get('priority', 3)}")
                    
                    if task.get('category'):
                        details.append(f"<b>Catégorie:</b> {task['category']}")
                    
                    if task.get('estimated_duration'):
                        details.append(f"<b>Durée estimée:</b> {task['estimated_duration']} min")
                    
                    if task.get('created_at'):
                        details.append(f"<b>Créée le:</b> {task['created_at']}")
                    
                    for detail in details:
                        story.append(Paragraph(detail, body_style))
                    
                    story.append(Spacer(1, 15))
            
            else:  # Simple template
                # Tasks table
                table_data = [['#', 'Titre', 'Statut', 'Priorité', 'Catégorie']]
                
                for i, task in enumerate(tasks, 1):
                    row = [
                        str(i),
                        task.get('title', 'Sans titre')[:40],
                        task.get('status', 'pending'),
                        str(task.get('priority', 3)),
                        task.get('category', '')[:20]
                    ]
                    table_data.append(row)
                
                tasks_table = Table(table_data, colWidths=[0.5*inch, 3*inch, 1*inch, 1*inch, 1.5*inch])
                tasks_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), self.header_color),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 9)
                ]))
                
                story.append(tasks_table)
            
            # Build PDF
            doc.build(story)
            
            self.logger.info(f"Tasks exported to PDF: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export tasks to PDF: {e}")
            return None
    
    def _export_tasks_docx(self, tasks: List[Dict], output_path: Path, 
                          template: ExportTemplate) -> Optional[Path]:
        """Export tasks to DOCX format."""
        if not Document:
            raise ImportError("python-docx is required for DOCX export")
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Rapport des Tâches - Easy Genie', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Generation info
            gen_para = doc.add_paragraph(
                f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            )
            gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Empty line
            
            # Summary
            doc.add_heading('Résumé', level=1)
            
            total_tasks = len(tasks)
            completed_tasks = len([t for t in tasks if t.get('status') == 'completed'])
            pending_tasks = total_tasks - completed_tasks
            
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Table Grid'
            
            summary_data = [
                ('Total des tâches', str(total_tasks)),
                ('Tâches terminées', str(completed_tasks)),
                ('Tâches en cours', str(pending_tasks)),
                ('Taux de completion', f"{(completed_tasks/total_tasks*100):.1f}%" if total_tasks > 0 else "0%")
            ]
            
            for i, (label, value) in enumerate(summary_data):
                summary_table.cell(i, 0).text = label
                summary_table.cell(i, 1).text = value
            
            doc.add_paragraph()  # Empty line
            
            # Tasks section
            doc.add_heading('Détail des Tâches', level=1)
            
            if template == ExportTemplate.DETAILED:
                for i, task in enumerate(tasks, 1):
                    # Task heading
                    task_heading = doc.add_heading(f"{i}. {task.get('title', 'Sans titre')}", level=2)
                    
                    # Task details
                    if task.get('description'):
                        desc_para = doc.add_paragraph()
                        desc_para.add_run('Description: ').bold = True
                        desc_para.add_run(task['description'])
                    
                    details_para = doc.add_paragraph()
                    details_para.add_run('Statut: ').bold = True
                    details_para.add_run(f"{task.get('status', 'pending')}\n")
                    details_para.add_run('Priorité: ').bold = True
                    details_para.add_run(f"{task.get('priority', 3)}\n")
                    
                    if task.get('category'):
                        details_para.add_run('Catégorie: ').bold = True
                        details_para.add_run(f"{task['category']}\n")
                    
                    if task.get('estimated_duration'):
                        details_para.add_run('Durée estimée: ').bold = True
                        details_para.add_run(f"{task['estimated_duration']} minutes\n")
                    
                    if task.get('created_at'):
                        details_para.add_run('Créée le: ').bold = True
                        details_para.add_run(task['created_at'])
            
            else:  # Simple template
                # Tasks table
                tasks_table = doc.add_table(rows=len(tasks) + 1, cols=5)
                tasks_table.style = 'Table Grid'
                
                # Header
                header_cells = tasks_table.rows[0].cells
                headers = ['#', 'Titre', 'Statut', 'Priorité', 'Catégorie']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].bold = True
                
                # Data rows
                for i, task in enumerate(tasks, 1):
                    row_cells = tasks_table.rows[i].cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = task.get('title', 'Sans titre')
                    row_cells[2].text = task.get('status', 'pending')
                    row_cells[3].text = str(task.get('priority', 3))
                    row_cells[4].text = task.get('category', '')
            
            # Save document
            doc.save(str(output_path))
            
            self.logger.info(f"Tasks exported to DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export tasks to DOCX: {e}")
            return None
    
    def _export_tasks_txt(self, tasks: List[Dict], output_path: Path, 
                         template: ExportTemplate) -> Optional[Path]:
        """Export tasks to TXT format."""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # Header
                f.write("RAPPORT DES TÂCHES - EASY GENIE\n")
                f.write("=" * 40 + "\n\n")
                f.write(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}\n\n")
                
                # Summary
                total_tasks = len(tasks)
                completed_tasks = len([t for t in tasks if t.get('status') == 'completed'])
                pending_tasks = total_tasks - completed_tasks
                
                f.write("RÉSUMÉ\n")
                f.write("-" * 20 + "\n")
                f.write(f"Total des tâches: {total_tasks}\n")
                f.write(f"Tâches terminées: {completed_tasks}\n")
                f.write(f"Tâches en cours: {pending_tasks}\n")
                if total_tasks > 0:
                    f.write(f"Taux de completion: {(completed_tasks/total_tasks*100):.1f}%\n")
                f.write("\n")
                
                # Tasks
                f.write("DÉTAIL DES TÂCHES\n")
                f.write("-" * 20 + "\n\n")
                
                for i, task in enumerate(tasks, 1):
                    f.write(f"{i}. {task.get('title', 'Sans titre')}\n")
                    
                    if template == ExportTemplate.DETAILED:
                        if task.get('description'):
                            f.write(f"   Description: {task['description']}\n")
                        f.write(f"   Statut: {task.get('status', 'pending')}\n")
                        f.write(f"   Priorité: {task.get('priority', 3)}\n")
                        if task.get('category'):
                            f.write(f"   Catégorie: {task['category']}\n")
                        if task.get('estimated_duration'):
                            f.write(f"   Durée estimée: {task['estimated_duration']} min\n")
                        if task.get('created_at'):
                            f.write(f"   Créée le: {task['created_at']}\n")
                    else:
                        f.write(f"   [{task.get('status', 'pending')}] Priorité: {task.get('priority', 3)}\n")
                    
                    f.write("\n")
            
            self.logger.info(f"Tasks exported to TXT: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export tasks to TXT: {e}")
            return None
    
    def _export_tasks_json(self, tasks: List[Dict], output_path: Path) -> Optional[Path]:
        """Export tasks to JSON format."""
        try:
            export_data = {
                'metadata': {
                    'export_date': datetime.now().isoformat(),
                    'total_tasks': len(tasks),
                    'completed_tasks': len([t for t in tasks if t.get('status') == 'completed']),
                    'export_format': 'json',
                    'application': 'Easy Genie Desktop'
                },
                'tasks': tasks
            }
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False, default=str)
            
            self.logger.info(f"Tasks exported to JSON: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export tasks to JSON: {e}")
            return None
    
    def _export_tasks_csv(self, tasks: List[Dict], output_path: Path) -> Optional[Path]:
        """Export tasks to CSV format."""
        try:
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                if not tasks:
                    return output_path
                
                # Define columns
                columns = ['id', 'title', 'description', 'status', 'priority', 
                          'category', 'estimated_duration', 'actual_duration', 
                          'created_at', 'updated_at', 'completed_at']
                
                writer = csv.DictWriter(f, fieldnames=columns)
                writer.writeheader()
                
                for task in tasks:
                    # Prepare row data
                    row = {col: task.get(col, '') for col in columns}
                    # Convert lists/dicts to strings
                    for key, value in row.items():
                        if isinstance(value, (list, dict)):
                            row[key] = json.dumps(value, ensure_ascii=False)
                    writer.writerow(row)
            
            self.logger.info(f"Tasks exported to CSV: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export tasks to CSV: {e}")
            return None
    
    def _export_tasks_html(self, tasks: List[Dict], output_path: Path, 
                          template: ExportTemplate) -> Optional[Path]:
        """Export tasks to HTML format."""
        try:
            html_content = f"""
<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Rapport des Tâches - Easy Genie</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        .header {{
            text-align: center;
            border-bottom: 2px solid #2E86AB;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .summary {{
            background-color: #f5f5f5;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 30px;
        }}
        .task {{
            border: 1px solid #ddd;
            margin-bottom: 15px;
            padding: 15px;
            border-radius: 5px;
        }}
        .task-title {{
            font-weight: bold;
            color: #2E86AB;
            margin-bottom: 10px;
        }}
        .task-meta {{
            color: #666;
            font-size: 0.9em;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #2E86AB;
            color: white;
        }}
        .status-completed {{ color: green; font-weight: bold; }}
        .status-pending {{ color: orange; font-weight: bold; }}
        .status-in-progress {{ color: blue; font-weight: bold; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Rapport des Tâches - Easy Genie</h1>
        <p>Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</p>
    </div>
    
    <div class="summary">
        <h2>Résumé</h2>
        <p><strong>Total des tâches:</strong> {len(tasks)}</p>
        <p><strong>Tâches terminées:</strong> {len([t for t in tasks if t.get('status') == 'completed'])}</p>
        <p><strong>Tâches en cours:</strong> {len(tasks) - len([t for t in tasks if t.get('status') == 'completed'])}</p>
    </div>
    
    <h2>Détail des Tâches</h2>
"""
            
            if template == ExportTemplate.DETAILED:
                for i, task in enumerate(tasks, 1):
                    status_class = f"status-{task.get('status', 'pending').replace('_', '-')}"
                    html_content += f"""
    <div class="task">
        <div class="task-title">{i}. {task.get('title', 'Sans titre')}</div>
        {f'<p><strong>Description:</strong> {task["description"]}</p>' if task.get('description') else ''}
        <div class="task-meta">
            <span class="{status_class}">Statut: {task.get('status', 'pending')}</span> | 
            Priorité: {task.get('priority', 3)} | 
            {f'Catégorie: {task["category"]} | ' if task.get('category') else ''}
            {f'Durée estimée: {task["estimated_duration"]} min | ' if task.get('estimated_duration') else ''}
            Créée le: {task.get('created_at', 'N/A')}
        </div>
    </div>
"""
            else:
                html_content += """
    <table>
        <thead>
            <tr>
                <th>#</th>
                <th>Titre</th>
                <th>Statut</th>
                <th>Priorité</th>
                <th>Catégorie</th>
                <th>Créée le</th>
            </tr>
        </thead>
        <tbody>
"""
                for i, task in enumerate(tasks, 1):
                    status_class = f"status-{task.get('status', 'pending').replace('_', '-')}"
                    html_content += f"""
            <tr>
                <td>{i}</td>
                <td>{task.get('title', 'Sans titre')}</td>
                <td class="{status_class}">{task.get('status', 'pending')}</td>
                <td>{task.get('priority', 3)}</td>
                <td>{task.get('category', '')}</td>
                <td>{task.get('created_at', 'N/A')}</td>
            </tr>
"""
                html_content += """
        </tbody>
    </table>
"""
            
            html_content += """
</body>
</html>
"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.logger.info(f"Tasks exported to HTML: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export tasks to HTML: {e}")
            return None
    
    def _export_tasks_markdown(self, tasks: List[Dict], output_path: Path, 
                              template: ExportTemplate) -> Optional[Path]:
        """Export tasks to Markdown format."""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # Header
                f.write("# Rapport des Tâches - Easy Genie\n\n")
                f.write(f"*Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}*\n\n")
                
                # Summary
                total_tasks = len(tasks)
                completed_tasks = len([t for t in tasks if t.get('status') == 'completed'])
                pending_tasks = total_tasks - completed_tasks
                
                f.write("## Résumé\n\n")
                f.write(f"- **Total des tâches:** {total_tasks}\n")
                f.write(f"- **Tâches terminées:** {completed_tasks}\n")
                f.write(f"- **Tâches en cours:** {pending_tasks}\n")
                if total_tasks > 0:
                    f.write(f"- **Taux de completion:** {(completed_tasks/total_tasks*100):.1f}%\n")
                f.write("\n")
                
                # Tasks
                f.write("## Détail des Tâches\n\n")
                
                if template == ExportTemplate.DETAILED:
                    for i, task in enumerate(tasks, 1):
                        f.write(f"### {i}. {task.get('title', 'Sans titre')}\n\n")
                        
                        if task.get('description'):
                            f.write(f"**Description:** {task['description']}\n\n")
                        
                        f.write(f"- **Statut:** {task.get('status', 'pending')}\n")
                        f.write(f"- **Priorité:** {task.get('priority', 3)}\n")
                        
                        if task.get('category'):
                            f.write(f"- **Catégorie:** {task['category']}\n")
                        
                        if task.get('estimated_duration'):
                            f.write(f"- **Durée estimée:** {task['estimated_duration']} minutes\n")
                        
                        if task.get('created_at'):
                            f.write(f"- **Créée le:** {task['created_at']}\n")
                        
                        f.write("\n---\n\n")
                
                else:  # Simple template
                    f.write("| # | Titre | Statut | Priorité | Catégorie | Créée le |\n")
                    f.write("|---|-------|--------|----------|-----------|----------|\n")
                    
                    for i, task in enumerate(tasks, 1):
                        title = task.get('title', 'Sans titre').replace('|', '\\|')
                        status = task.get('status', 'pending')
                        priority = str(task.get('priority', 3))
                        category = task.get('category', '').replace('|', '\\|')
                        created = task.get('created_at', 'N/A')
                        
                        f.write(f"| {i} | {title} | {status} | {priority} | {category} | {created} |\n")
            
            self.logger.info(f"Tasks exported to Markdown: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export tasks to Markdown: {e}")
            return None
    
    # Brain dump export methods
    def export_brain_dump(self, brain_dump: Dict, format_type: ExportFormat = None, 
                         filename: str = None) -> Optional[Path]:
        """Export a brain dump to specified format."""
        format_type = format_type or self.default_format
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            title_safe = brain_dump.get('title', 'brain_dump').replace(' ', '_')[:20]
            filename = f"{title_safe}_{timestamp}.{format_type.value}"
        
        output_path = self.output_directory / filename
        
        try:
            if format_type == ExportFormat.TXT:
                return self._export_brain_dump_txt(brain_dump, output_path)
            elif format_type == ExportFormat.JSON:
                return self._export_brain_dump_json(brain_dump, output_path)
            elif format_type == ExportFormat.MD:
                return self._export_brain_dump_markdown(brain_dump, output_path)
            else:
                # For other formats, convert to tasks-like structure
                pseudo_tasks = [{
                    'title': brain_dump.get('title', 'Décharge de pensées'),
                    'description': brain_dump.get('content', ''),
                    'status': 'completed',
                    'priority': 1,
                    'category': 'Brain Dump',
                    'created_at': brain_dump.get('created_at', ''),
                    'word_count': brain_dump.get('word_count', 0),
                    'character_count': brain_dump.get('character_count', 0)
                }]
                return self.export_tasks(pseudo_tasks, format_type, ExportTemplate.DETAILED, filename)
                
        except Exception as e:
            self.logger.error(f"Failed to export brain dump: {e}")
            return None
    
    def _export_brain_dump_txt(self, brain_dump: Dict, output_path: Path) -> Optional[Path]:
        """Export brain dump to TXT format."""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("DÉCHARGE DE PENSÉES - EASY GENIE\n")
                f.write("=" * 40 + "\n\n")
                
                if brain_dump.get('title'):
                    f.write(f"Titre: {brain_dump['title']}\n\n")
                
                f.write(f"Créé le: {brain_dump.get('created_at', 'N/A')}\n")
                f.write(f"Nombre de mots: {brain_dump.get('word_count', 0)}\n")
                f.write(f"Nombre de caractères: {brain_dump.get('character_count', 0)}\n\n")
                
                f.write("CONTENU\n")
                f.write("-" * 20 + "\n\n")
                f.write(brain_dump.get('content', ''))
                
                # Add analysis if available
                analysis_data = brain_dump.get('analysis_data', {})
                if analysis_data and isinstance(analysis_data, dict):
                    f.write("\n\n" + "=" * 40 + "\n")
                    f.write("ANALYSE IA\n")
                    f.write("-" * 20 + "\n\n")
                    for key, value in analysis_data.items():
                        f.write(f"{key.title()}: {value}\n\n")
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export brain dump to TXT: {e}")
            return None
    
    def _export_brain_dump_json(self, brain_dump: Dict, output_path: Path) -> Optional[Path]:
        """Export brain dump to JSON format."""
        try:
            export_data = {
                'metadata': {
                    'export_date': datetime.now().isoformat(),
                    'export_format': 'json',
                    'application': 'Easy Genie Desktop',
                    'content_type': 'brain_dump'
                },
                'brain_dump': brain_dump
            }
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False, default=str)
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export brain dump to JSON: {e}")
            return None
    
    def _export_brain_dump_markdown(self, brain_dump: Dict, output_path: Path) -> Optional[Path]:
        """Export brain dump to Markdown format."""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("# Décharge de Pensées - Easy Genie\n\n")
                
                if brain_dump.get('title'):
                    f.write(f"## {brain_dump['title']}\n\n")
                
                f.write(f"*Créé le {brain_dump.get('created_at', 'N/A')}*\n\n")
                
                # Metadata
                f.write("### Statistiques\n\n")
                f.write(f"- **Nombre de mots:** {brain_dump.get('word_count', 0)}\n")
                f.write(f"- **Nombre de caractères:** {brain_dump.get('character_count', 0)}\n\n")
                
                # Content
                f.write("### Contenu\n\n")
                f.write(brain_dump.get('content', ''))
                
                # Analysis
                analysis_data = brain_dump.get('analysis_data', {})
                if analysis_data and isinstance(analysis_data, dict):
                    f.write("\n\n### Analyse IA\n\n")
                    for key, value in analysis_data.items():
                        f.write(f"**{key.title()}:**\n\n{value}\n\n")
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to export brain dump to Markdown: {e}")
            return None
    
    # Configuration methods
    def set_output_directory(self, directory: Union[str, Path]):
        """Set the output directory for exports."""
        self.output_directory = Path(directory)
        self.output_directory.mkdir(parents=True, exist_ok=True)
        
        if self.settings_manager:
            self.settings_manager.set('export.output_directory', str(self.output_directory))
        
        self.logger.info(f"Export directory set to: {self.output_directory}")
    
    def set_default_format(self, format_type: ExportFormat):
        """Set the default export format."""
        self.default_format = format_type
        
        if self.settings_manager:
            self.settings_manager.set('export.default_format', format_type.value)
        
        self.logger.info(f"Default export format set to: {format_type.value}")
    
    def get_supported_formats(self) -> List[ExportFormat]:
        """Get list of supported export formats."""
        supported = [ExportFormat.TXT, ExportFormat.JSON, ExportFormat.CSV, 
                    ExportFormat.HTML, ExportFormat.MD]
        
        if reportlab:
            supported.append(ExportFormat.PDF)
        
        if Document:
            supported.append(ExportFormat.DOCX)
        
        return supported
    
    def get_status(self) -> Dict:
        """Get export service status."""
        return {
            'default_format': self.default_format.value,
            'default_template': self.default_template.value,
            'output_directory': str(self.output_directory),
            'supported_formats': [f.value for f in self.get_supported_formats()],
            'pdf_available': reportlab is not None,
            'docx_available': Document is not None,
            'output_directory_exists': self.output_directory.exists(),
            'output_directory_writable': self.output_directory.exists() and 
                                       self.output_directory.is_dir()
        }